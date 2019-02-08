from django.shortcuts import render
from learn.models import *
from django.urls import reverse
from django.http import HttpResponse
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from bs4 import BeautifulSoup
from django.conf import settings
from django.db.models import Q
from datetime import datetime
from itertools import  chain
from django.core.mail import send_mail,BadHeaderError
import time
import urllib.request
import docx
import operator
import re
import os
# Create your views here.


def index(request):
    return render(request,'登录.html')


def openCrawl(request):
    return render(request,'知识获取.html')


def login(request):
    username = request.POST.get('loginname', '')
    password = request.POST.get('password', '')
    if len(password)<8:
        string=u"密码的长度应大于8位"
        return render(request, '登录.html', {'string': string})
    else:
     if Worker.objects.filter(userNo=username).filter(pwd=password):
       workeritem=Worker.objects.get(userNo=username)
       name=str(workeritem.userName)
       settings.USER_NAME=name
       greet="hello!"+name
       return render(request, '主页.html', {'string': greet})
     else:
        string=u"您的账号或是密码可能不正确"
        return render(request, '登录.html', {'string': string})


def openUpload(request):
     localtime=time.strftime('%Y-%m-%d', time.localtime(time.time()))
     return render(request, '内部知识上传.html',{'username':settings.USER_NAME,'localtime':localtime})


def upload(request):
     knowledgetype = request.POST.get('knowledgeType')
     localtime = time.strftime('%Y-%m-%d', time.localtime(time.time()))
     if knowledgetype == 'file':
       knowledgeitem = request.FILES.get('fileitem')
       #path=default_storage.save('D:/system/mysystem/learn/knowledge/file/',ContentFile(knowledgeitem.read()))
       knowledgeTitle = knowledgeitem.name
       knowledgeField = request.POST.get('knowledgeField')
       knowledgefileNo = request.POST.get('knowledgeModelNo')
       knowledgedevotionTime = request.POST.get('knowledgedevotionTime')
       knowledgedevoter = request.POST.get('knowledgedevoter')
       knowledgeRemarks = request.POST.get('knowledgeRemarks')
       if File.objects.filter(title__icontains=knowledgeTitle).exists():
           itemcount=File.objects.filter(title__icontains=knowledgeTitle).count()
           count=str(itemcount+1)
           knowledgefileNo=knowledgefileNo+'-'+count
       else:
           knowledgefileNo=knowledgefileNo+'-'+'1'
       if knowledgedevotionTime is '':
          knowledgedevotionTime = time.strftime('%Y-%m-%d', time.localtime(time.time()))
       File.objects.get_or_create(title=knowledgeTitle, type=knowledgetype, field=knowledgeField, devotionTime=str(knowledgedevotionTime), fileNo=knowledgefileNo, devoter=knowledgedevoter, remarks=knowledgeRemarks, link=knowledgeitem)
       return render(request,'内部知识上传.html',{'username':settings.USER_NAME,'localtime':localtime})
     if knowledgetype == 'chart':
        knowledgeitem = request.FILES.get('fileitem')
        path = default_storage.save('D:/system/mysystem/learn/knowledge/chart/', ContentFile(knowledgeitem.read()))
        knowledgeTitle = knowledgeitem.name
        knowledgeField = request.POST.get('knowledgeField')
        knowledgechartNo = request.POST.get('knowledgeModelNo')
        knowledgedevotionTime = request.POST.get('knowledgedevotionTime')
        knowledgedevoter = request.POST.get('knowledgedevoter')
        knowledgeRemarks = request.POST.get('knowledgeRemarks')
        if knowledgedevotionTime is '':
            knowledgedevotionTime = time.strftime('%Y-%m-%d', time.localtime(time.time()))
        Chart.objects.get_or_create(title=knowledgeTitle, type=knowledgetype, field=knowledgeField,
                                   devotionTime=str(knowledgedevotionTime), chartNo=knowledgechartNo,
                                   devoter=knowledgedevoter, remarks=knowledgeRemarks, link=knowledgeitem)
        return render(request, '内部知识上传.html',{'username':settings.USER_NAME,'localtime':localtime})
     if knowledgetype == 'data':
        knowledgeitem = request.FILES.get('fileitem')
        #path = default_storage.save('D:/system/mysystem/learn/knowledge/data/', ContentFile(knowledgeitem.read()))
        knowledgeTitle = knowledgeitem.name
        knowledgeField = request.POST.get('knowledgeField')
        knowledgedataNo = request.POST.get('knowledgeModelNo')
        knowledgedevotionTime = request.POST.get('knowledgedevotionTime')
        knowledgedevoter = request.POST.get('knowledgedevoter')
        knowledgeRemarks = request.POST.get('knowledgeRemarks')
        if knowledgedevotionTime is '':
            knowledgedevotionTime = time.strftime('%Y-%m-%d', time.localtime(time.time()))
        Data.objects.get_or_create(title=knowledgeTitle, type=knowledgetype, field=knowledgeField,
                                   devotionTime=str(knowledgedevotionTime), dataNo=knowledgedataNo,
                                   devoter=knowledgedevoter, remarks=knowledgeRemarks, link=knowledgeitem)
        return render(request, '内部知识上传.html',{'username':settings.USER_NAME,'localtime':localtime})


def openQuery(request):
    return render(request,'知识查询.html')


def queryKnowledge(request):
       queryStr=''
       knowledgeType = request.POST.get('knowledgeType')
       knowledgeNo = request.POST.get('queryKnowledgeNo')
       knowledgeField = request.POST.get('queryKnowledgeField')
       startTime = request.POST.get('startTime')
       endTime = request.POST.get('endTime')
       startTime = datetime.strptime(startTime, '%Y-%m-%d')
       endTime = datetime.strptime(endTime, '%Y-%m-%d')
       if knowledgeType == 'file':
          if knowledgeNo == '':
             if File.objects.filter(Q(field=knowledgeField)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                itemcount = File.objects.filter(Q(field=knowledgeField)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).count()
                knowledgeList = File.objects.filter((Q(field=knowledgeField)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)))
                return render(request, '知识查询.html', {'dataList': knowledgeList,'warnstring':int(itemcount)})
          if knowledgeField == '':
             if File.objects.filter(Q(fileNo=knowledgeNo)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                 knowledgeList = File.objects.filter(Q(fileNo=knowledgeNo)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime))
                 return render(request, '知识查询.html', {'dataList': knowledgeList})
          if  knowledgeField == '' and knowledgeNo == '':
              if File.objects.filter(Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                 itemcount = File.objects.filter(Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).count()
                 knowledgeList = File.objects.filter(Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime))
                 return render(request, '知识查询.html', {'dataList': knowledgeList,'warnstring':int(itemcount)})
          else:
              if File.objects.filter(Q(fileNo=knowledgeNo)&Q(field=knowledgeField)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                 knowledgeList = File.objects.filter(Q(fileNo=knowledgeNo)&Q(field=knowledgeField)&Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime))
                 return render(request, '知识查询.html', {'dataList': knowledgeList})
       if knowledgeType=='data':
           if knowledgeNo == '':
               if Data.objects.filter(Q(field=knowledgeField) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter((Q(field=knowledgeField) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if knowledgeField == '':
               if Data.objects.filter(Q(dataNo=knowledgeNo) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter(Q(fileNo=knowledgeNo) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if knowledgeField == '' and knowledgeNo == '':
               if Data.objects.filter(Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter(Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if Data.objects.filter(Q(dataNo=knowledgeNo) &Q(field=knowledgeField) | Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = Data.objects.filter(Q(dataNo=knowledgeNo) & Q(field=knowledgeField) | Q(devotionTime__gte=startTime)& Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
       if knowledgeType == 'chart':
           if knowledgeNo == '':
               if Chart.objects.filter(Q(field=knowledgeField) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter((Q(field=knowledgeField) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if knowledgeField == '':
               if Chart.objects.filter(Q(chartNo=knowledgeNo) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter(Q(fileNo=knowledgeNo) & Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if knowledgeField == '' and knowledgeNo == '':
               if Chart.objects.filter(Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = File.objects.filter(Q(devotionTime__gte=startTime) & Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
           if Chart.objects.filter(Q(chartNo=knowledgeNo) &Q(field=knowledgeField) |Q(devotionTime__gte=startTime)&Q(devotionTime__lte=endTime)).exists():
                   knowledgeList = Chart.objects.filter(Q(chartNo=knowledgeNo) & Q(field=knowledgeField)| Q(devotionTime__gte=startTime)& Q(devotionTime__lte=endTime))
                   return render(request, '知识查询.html', {'dataList': knowledgeList})
       if knowledgeType == 'all':
           fileList = File.objects.filter(Q(fileNo=knowledgeNo) &Q(field=knowledgeField) | Q(devotionTime__gte=startTime) | Q(devotionTime__lte=endTime))
           dataList = Data.objects.filter(Q(dataNo=knowledgeNo) &Q(field=knowledgeField) | Q(devotionTime__gte=startTime) | Q(devotionTime__lte=endTime))
           chartList = Chart.objects.filter(Q(chartNo=knowledgeNo) & Q(field=knowledgeField) | Q(devotionTime__gte=startTime) | Q(devotionTime__lte=endTime))
           knowledgeList = chain(fileList,dataList,chartList)
           return render(request,'知识查询.html',{'dataList':knowledgeList})
       string=u"没有查询到对象!"
       return render(request, '知识查询.html', {'warningstring': string})


def crawlKnowledge(request):
    crawllink = request.POST.get('knowledgeLink')
    html = urllib.request.urlopen(crawllink).read()
    soup = BeautifulSoup(html,'html.parser')
    [s.extract() for s in soup('script')]
    text = soup.get_text()
    text = text.strip()
    text = text.replace('\n', '')
    title = soup.title.get_text()
    return render(request, '知识获取.html', {'crawlTitle': text, 'textTitle': title})


def sendemail(request):
    workerlist=Worker.objects.all()
    return render(request, '邮件发送.html',{'workerlist':workerlist})


def email(request):
    workerlist = Worker.objects.all()
    toemail = request.POST.get('toemail')
    subject = request.POST.get('subject')
    content = request.POST.get('content')
    try:
        send_mail(subject, content, 'yifzhou@163.com', [toemail])
    except BadHeaderError:
        return HttpResponse('Invalid header found.')
    return render(request, '邮件发送.html',{'workerlist':workerlist})


def openCompareKnowledge(request):
    return render(request, '知识比对.html')


def getContent(request):
    knowledgeNo=request.POST.get('knowledgeNo')
    knowledgeversion=request.POST.get('knowledgeversion')
    itemcount=File.objects.filter(fileNo__icontains=knowledgeNo).count()
    itemcount=str(itemcount)
    newknowledgeNo=knowledgeNo+'-'+itemcount
    oldknowledgeNo=knowledgeNo+'-'+knowledgeversion
    newfileitem=File.objects.get(fileNo=newknowledgeNo)
    fileitem=File.objects.get(fileNo=oldknowledgeNo)
    newtitle=newfileitem.title
    newpath=newfileitem.link.path
    oldtitle=fileitem.title
    oldpath=fileitem.link.path
    doc1 = docx.Document(newpath)
    doc2=docx.Document(oldpath)
    ttext1 = ''
    ttext2=''
    for paragtaph in doc1.paragraphs:
        text = paragtaph.text
        ttext1 = ttext1 + text
    for paragtaph in doc2.paragraphs:
        text = paragtaph.text
        ttext2 = ttext2 + text
    return render(request, '知识比对.html', {'content1': ttext1, 'content2': ttext2})


def processstr(example):
    example=re.split(r'[，。]', example)
    dict={}
    linenum=1
    for item in example:
        dict[linenum] = item
        linenum = linenum+1
    return dict


def com(dict1, dict2):
    list = []
    len1 = len(dict1)
    len2 = len(dict2)
    if len1 == len2:
        for i in range(len1):
            if operator.eq(dict1.get(i+1), dict2.get(i+1)) == False:
                list.append(i+1)
    elif len1 > len2:
        for i in range(len2):
            if operator.eq(dict1.get(i+1), dict2.get(i+1)) == False:
                list.append(i+1)
        last = i+1
        for i in range(last+1, len1+1):
            list.append(i)
    else:
        for i in range(len1):
            if operator.eq(dict1.get(i+1), dict2.get(i+1)) == False:
                list.append(i+1)
        last = i+1
        for i in range(last+1, len2+1):
            list.append(i)
    return list


def compareContent(request):
    finaldict = {}
    content1 = request.POST.get('content1')
    content2 = request.POST.get('content2')
    dict1 = processstr(content1)
    dict2 = processstr(content2)
    list = com(dict1, dict2)
    for key in list:
        finaldict[key]=(dict1.get(key), dict2.get(key))
    return render(request, '知识比对.html', {'content3': finaldict})











